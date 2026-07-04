"""Tests for the upload-pipeline (backend/app/upload.py + extract wrappers).

  * test_extract_wrappers_mock : unit test of the in-memory extraction wrappers
    with the deterministic mock LLM provider (no services, no cost).
  * test_upload_endpoint_docx  : full POST /api/upload -> done flow via TestClient
    with the mock LLM, asserting stages, searchability and a valid graph fragment.
    Skipped when Neo4j / Elasticsearch are not reachable.
"""
import io
import os
import sys
import time
from pathlib import Path

import pytest

BACKEND = Path(__file__).resolve().parent.parent
REPO_ROOT = BACKEND.parent
for p in (str(BACKEND), str(REPO_ROOT)):
    if p not in sys.path:
        sys.path.insert(0, p)

from tests.conftest import NEO4J_AVAILABLE  # noqa: E402


def _mk_docx(path: Path) -> None:
    from docx import Document
    d = Document()
    d.add_heading("Тестовый эксперимент по извлечению никеля", 0)
    d.add_paragraph(
        "Синтетический электролит содержал 85 г/л Ni2+. Опыты вели при температуре "
        "65 °C и плотности тока 300 А/м2; извлечение никеля составило 87,5 %.")
    d.add_paragraph(
        "Оптимальным режимом признана температура 65 °C, обеспечивающая извлечение "
        "никеля 87,5 % при поддержании pH 3.")
    d.save(path)


def test_extract_wrappers_mock():
    """extract_payloads + build_fragment work end-to-end on the mock provider."""
    os.environ["LLM_PROVIDER"] = "mock"
    from pipeline.extract import runner

    chunks = [
        {"chunk_id": "t_c0000", "doc_id": "t_doc", "section_title": "Результаты",
         "text": "Извлечение никеля составило 87,5 % при температуре 65 °C."},
        {"chunk_id": "t_c0001", "doc_id": "t_doc", "section_title": None,
         "text": "Электролит содержал 85 г/л Ni2+ при плотности тока 300 А/м2."},
    ]
    rows = runner.extract_payloads(chunks, model="lite")
    assert len(rows) == 2
    assert all("payload" in r and "ok" in r for r in rows)

    gb = runner.build_fragment(rows, use_embedding=False)
    # a Publication node for the doc must always be present after processing chunks
    assert "pub:t_doc" in gb.nodes
    assert gb.nodes["pub:t_doc"]["type"] == "Publication"
    # nodes/edges are plain dicts ready for the loader contract
    for n in gb.nodes.values():
        assert {"id", "type", "name"} <= set(n)


@pytest.mark.skipif(not NEO4J_AVAILABLE, reason="Neo4j not reachable (start docker compose)")
def test_upload_endpoint_docx(tmp_path):
    os.environ["LLM_PROVIDER"] = "mock"
    from app import db as appdb
    if not appdb.es_available():
        pytest.skip("Elasticsearch not reachable")

    from fastapi.testclient import TestClient
    from app.main import app

    docx_path = tmp_path / "upload_test_experiment.docx"
    _mk_docx(docx_path)

    with TestClient(app) as client:
        with open(docx_path, "rb") as fh:
            r = client.post("/api/upload", files={"file": (docx_path.name, fh)})
        assert r.status_code == 200, r.text
        body = r.json()
        job_id, doc_id = body["job_id"], body["doc_id"]
        assert doc_id.startswith("up_")

        # poll to completion
        job = None
        for _ in range(180):
            job = client.get(f"/api/upload/{job_id}").json()
            if job["stage"] in ("done", "failed"):
                break
            time.sleep(0.5)
        assert job["stage"] == "done", job

        res = job["result"]
        assert res["n_chunks"] >= 1
        # every canonical stage ran and is reported
        stages = res["stages"]
        for st in ("extracting_text", "chunking", "embedding", "indexing",
                   "extracting_knowledge", "merging_graph"):
            assert st in stages, st
        # graph fragment is valid and bounded
        gp = res["graph_preview"]
        assert len(gp["nodes"]) <= 30
        assert any(n["id"] == f"pub:{doc_id}" for n in gp["nodes"])
        edge_ids = {n["id"] for n in gp["nodes"]}
        for e in gp["edges"]:
            assert e["src"] in edge_ids and e["dst"] in edge_ids

        # document is immediately searchable via full-text
        sr = client.post("/api/search",
                         json={"query": "извлечение никеля 87,5 % при 65 °C"}).json()
        cited = [c.get("doc_id") for c in (sr.get("citations") or [])]
        assert doc_id in cited, cited

        # publication node landed in Neo4j
        assert appdb.get_node(f"pub:{doc_id}") is not None

        # dedup: same bytes -> same doc_id, cached
        with open(docx_path, "rb") as fh:
            r2 = client.post("/api/upload", files={"file": (docx_path.name, fh)}).json()
        assert r2["doc_id"] == doc_id
        assert r2["cached"] is True
