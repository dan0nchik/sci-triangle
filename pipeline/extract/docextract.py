"""
docextract.py — извлечение текста из DOCX/PDF и чанкование (dev-нужды направления B).

Это НЕ подмена направления A. Модуль используется только чтобы собрать
dev-выборку `corpus/dev_chunks.jsonl` из golden-документов, пока агент A
готовит полный корпус. Формат чанков совпадает с контрактом PLAN §4.1.
"""

from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Iterator


def extract_docx(path: str | Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    # таблицы -> строки через таб
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)


def extract_pdf(path: str | Path) -> str:
    import fitz  # PyMuPDF

    parts: list[str] = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            t = page.get_text("text").strip()
            if t:
                parts.append(t)
    return "\n".join(parts)


def extract_text(path: str | Path) -> str:
    p = Path(path)
    suf = p.suffix.lower()
    if suf in (".docx", ".docm"):
        return extract_docx(p)
    if suf == ".pdf":
        return extract_pdf(p)
    raise ValueError(f"неподдерживаемый формат для dev-извлечения: {suf}")


_WS = re.compile(r"[ \t ]+")


def _clean(text: str) -> str:
    text = text.replace("\r", "\n")
    text = _WS.sub(" ", text)
    # схлопываем множественные пустые строки
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _approx_tokens(s: str) -> int:
    # грубая оценка: ~4 символа/токен для смешанного RU/EN
    return max(1, len(s) // 4)


def chunk_text(
    text: str,
    target_tokens: int = 900,
    overlap_tokens: int = 100,
) -> list[str]:
    """Чанкование по абзацам с overlap. Возвращает список текстов чанков."""
    text = _clean(text)
    paras = [p.strip() for p in text.split("\n") if p.strip()]
    chunks: list[str] = []
    cur: list[str] = []
    cur_tok = 0
    for para in paras:
        pt = _approx_tokens(para)
        if cur and cur_tok + pt > target_tokens:
            chunks.append("\n".join(cur))
            # overlap: тянем хвостовые абзацы
            back: list[str] = []
            back_tok = 0
            for p in reversed(cur):
                back.insert(0, p)
                back_tok += _approx_tokens(p)
                if back_tok >= overlap_tokens:
                    break
            cur = back
            cur_tok = back_tok
        cur.append(para)
        cur_tok += pt
    if cur:
        chunks.append("\n".join(cur))
    # выкидываем совсем мусорные короткие
    return [c for c in chunks if len(c) > 40]


def doc_id_for(path: str | Path, prefix: str = "dev") -> str:
    h = hashlib.sha1(str(Path(path).name).encode("utf-8")).hexdigest()[:6]
    return f"{prefix}{h}"


def iter_chunks(path: str | Path, **kw) -> Iterator[dict]:
    """Извлекает текст файла и выдаёт чанки в формате corpus/chunks.jsonl (§4.1)."""
    text = extract_text(path)
    did = doc_id_for(path)
    lang = "ru" if re.search(r"[а-яА-Я]", text) else "en"
    for seq, ct in enumerate(chunk_text(text, **kw)):
        clang = "ru" if re.search(r"[а-яА-Я]", ct) else "en"
        yield {
            "chunk_id": f"{did}_c{seq:04d}",
            "doc_id": did,
            "seq": seq,
            "text": ct,
            "n_tokens": _approx_tokens(ct),
            "page_from": None,
            "page_to": None,
            "lang": clang or lang,
            "section_title": None,
        }
